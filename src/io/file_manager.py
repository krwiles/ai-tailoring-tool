import shutil
from datetime import date
from pathlib import Path
import re

from docx import Document
from typing import List

from src.config import AppSettings
from src.model import JobData


class FileManager:
    def __init__(self, settings: AppSettings) -> None:
        self.settings = settings
        self.output_dir = settings.output_dir
        self.project_dir = settings.project_dir

    def load_text(self, file_path: Path) -> str:
        """Read text file into string"""
        full_path = self.project_dir / file_path
        if not full_path.exists():
            raise FileNotFoundError(f"Expected text file at: {full_path}")
        return full_path.read_text(encoding="utf-8")

    def read_docx_paragraphs(self, file_path: Path) -> List[str]:
        """Read docx file into list of paragraphs"""
        full_path = self.project_dir / file_path
        if not full_path.exists():
            raise FileNotFoundError(f"Expected docx file at: {full_path}")
        doc = Document(full_path)
        return [para.text for para in doc.paragraphs]

    def read_docx(self, file_path: Path) -> str:
        """Read docx file into one string"""
        return "\n".join(self.read_docx_paragraphs(file_path))

    def write_docx(self, text: str, file_path: Path) -> None:
        """Write to docx"""
        doc = Document()
        for para_text in text.split("\n"):
            doc.add_paragraph(para_text)
        doc.save(self.output_dir / file_path)

    def copy_resume(self, destination: Path):
        """Copy the default resume to the destination"""
        resume = (self.project_dir / "data" / "resume.docx").resolve()
        if not resume.exists():
            raise FileNotFoundError(f"Resume not found: {resume}")

        self.create_directory(destination)
        dest_folder = (self.output_dir / destination).resolve()
        shutil.copy2(resume, dest_folder / self.settings.resume_name)

    def create_directory(self, dir_path: Path):
        full_path = self.output_dir / dir_path
        full_path.mkdir(parents=True, exist_ok=True)

    def copy_cover_letter(self, destination: Path, job_data: JobData):
        """Copy the default cover letter preserving formatting"""
        cover_letter = Document(self.project_dir / "data" / "cover_letter.docx")
        today = date.today().strftime("%B %d, %Y")

        for para in cover_letter.paragraphs:
            for run in para.runs:
                run.text = run.text.format(
                    date=today,
                    location=job_data.location,
                    company=job_data.company,
                    position=job_data.job_title
                )

        self.create_directory(destination)
        cover_letter.save(self.output_dir / destination / self.settings.cover_letter_name)

    def sanitize_dirname(self, name: str, replacement: str = "_") -> str:
        if not name:
            return "untitled"
        # Replace forbidden characters (Windows + cross-platform)
        name = re.sub(r'[<>:"/\\|?*]', replacement, name)
        # Collapse whitespace
        name = re.sub(r'\s+', ' ', name).strip()
        # Remove trailing dots/spaces (Windows issue)
        name = name.rstrip(" .")

        return name
